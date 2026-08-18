from __future__ import annotations

from pathlib import Path
from shutil import copy2

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path('/Users/apple/Desktop/recomend movie project')
OUT = ROOT / 'Musubi_발표자료/최신_2026-08-17/2026-08-18'
ASSET = OUT / 'assets'
OUT.mkdir(parents=True, exist_ok=True)
ASSET.mkdir(parents=True, exist_ok=True)

FONT_PATH = '/System/Library/Fonts/AppleSDGothicNeo.ttc'
FONT = 'Arial'
BLUE = '1F4E78'
NAVY = '17324D'
LIGHT = 'EAF2F8'
PALE = 'F5F7FA'
GRAY = '667085'
GREEN = '1F7A4D'
ORANGE = 'D97706'
RED = 'B42318'


def font(size: int, bold: bool = False):
    index = 8 if bold else 0
    try:
        return ImageFont.truetype(FONT_PATH, size, index=index)
    except TypeError:
        return ImageFont.truetype(FONT_PATH, size)


def rounded(draw, box, fill, outline='#CBD5E1', width=2, radius=18):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def centered(draw, box, text, fnt, fill='#172B4D'):
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=fnt, spacing=5, align='center')
    w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.multiline_text(((x1 + x2 - w) / 2, (y1 + y2 - h) / 2), text, font=fnt, fill=fill, spacing=5, align='center')


def arrow(draw, start, end, color='#64748B', width=5):
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    dx, dy = ex - sx, ey - sy
    mag = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / mag, dy / mag
    px, py = -uy, ux
    size = 14
    p1 = (ex, ey)
    p2 = (ex - ux * size + px * size * .55, ey - uy * size + py * size * .55)
    p3 = (ex - ux * size - px * size * .55, ey - uy * size - py * size * .55)
    draw.polygon([p1, p2, p3], fill=color)


def make_logical_diagram(path: Path):
    im = Image.new('RGB', (1800, 1050), '#F8FAFC')
    d = ImageDraw.Draw(im)
    d.text((70, 35), 'CineVerse 멀티 AZ 논리 구조', font=font(48, True), fill='#102A43')
    d.text((70, 95), '최종 확정 목표 · 관리형 LB + KKE + 전용 GPU/DB', font=font(25), fill='#52606D')

    boxes = [
        ((70, 180, 300, 280), '사용자\nmovieverse.cloud', '#FFFFFF'),
        ((375, 170, 680, 290), 'Public LB HA Group\nLB-A / LB-B', '#E8F1FB'),
        ((760, 150, 1140, 310), 'Private KKE Cluster\nWorker 4대 (AZ별 2대)\nFrontend / Backend Pods', '#EAF7EF'),
        ((1230, 155, 1700, 305), 'Internal AI LB HA Group\nAI LB-A / AI LB-B', '#FFF4E5'),
        ((1230, 390, 1700, 570), 'GPU AI VM 2대\nGPU-A / GPU-B\nAI FastAPI · llama-server\nMilvus · etcd · MinIO', '#FFF8ED'),
        ((760, 430, 1140, 570), 'PostgreSQL 전용 VM\nPrimary + PgBouncer', '#F2EDFF'),
        ((760, 690, 1140, 830), 'PostgreSQL Standby\nStreaming replication\n수동 승격', '#F7F3FF'),
        ((1230, 700, 1700, 830), 'Object Storage\nDB backup / assets', '#EEF2F6'),
        ((70, 700, 560, 830), 'Container Registry\nFrontend / Backend image', '#EEF2F6'),
    ]
    for box, text, fill in boxes:
        rounded(d, box, fill)
        centered(d, box, text, font(25, True if '\n' not in text else False))

    arrow(d, (300, 230), (375, 230))
    arrow(d, (680, 230), (760, 230))
    arrow(d, (1140, 230), (1230, 230))
    arrow(d, (1465, 305), (1465, 390))
    arrow(d, (950, 310), (950, 430))
    arrow(d, (950, 570), (950, 690))
    arrow(d, (1140, 500), (1230, 500))
    arrow(d, (1140, 760), (1230, 760))
    arrow(d, (560, 760), (760, 760))

    d.text((70, 910), '공개 포트: LB 80/443만 허용  ·  DB/GPU/Worker는 Private subnet  ·  SSH는 NAT+Bastion 경유', font=font(27, True), fill='#334E68')
    im.save(path)


def make_physical_diagram(path: Path):
    im = Image.new('RGB', (1800, 1120), '#F8FAFC')
    d = ImageDraw.Draw(im)
    d.text((70, 35), 'CineVerse 멀티 AZ 물리 구조', font=font(48, True), fill='#102A43')
    d.text((70, 95), 'VPC 10.30.0.0/16 · VM 총 10대', font=font(25), fill='#52606D')

    azs = [
        (70, 160, 865, 980, 'kr-central-2-a', 'Public-A 10.30.1.0/24', 'Private-A 10.30.2.0/24', ['NAT+Bastion-A'], ['KKE Worker-A1', 'KKE Worker-A2', 'GPU-A', 'PostgreSQL Primary']),
        (935, 160, 1730, 980, 'kr-central-2-b', 'Public-B 10.30.4.0/24', 'Private-B 10.30.3.0/24', ['NAT+Bastion-B'], ['KKE Worker-B1', 'KKE Worker-B2', 'GPU-B', 'PostgreSQL Standby']),
    ]
    for x1, y1, x2, y2, az, pub, priv, pubs, privs in azs:
        rounded(d, (x1, y1, x2, y2), '#FFFFFF', '#94A3B8', 3, 24)
        d.text((x1 + 30, y1 + 22), az, font=font(31, True), fill='#17324D')
        rounded(d, (x1 + 30, y1 + 85, x2 - 30, y1 + 270), '#FFF4E5', '#F59E0B', 3, 18)
        d.text((x1 + 55, y1 + 105), pub, font=font(25, True), fill='#92400E')
        centered(d, (x1 + 100, y1 + 160, x2 - 100, y1 + 245), pubs[0], font(27, True), '#7C2D12')
        rounded(d, (x1 + 30, y1 + 310, x2 - 30, y2 - 35), '#EAF7EF', '#22C55E', 3, 18)
        d.text((x1 + 55, y1 + 330), priv, font=font(25, True), fill='#166534')
        positions = [(x1+70, y1+400, x1+355, y1+515), (x1+405, y1+400, x2-70, y1+515), (x1+70, y1+565, x1+355, y1+680), (x1+405, y1+565, x2-70, y1+680)]
        for b, label in zip(positions, privs):
            rounded(d, b, '#FFFFFF', '#86B99A', 2, 14)
            centered(d, b, label, font(24, True))

    d.text((70, 1015), '관리형 자원(별도): Public LB HA · Internal AI LB HA · KKE control plane · Container Registry · Object Storage', font=font(27, True), fill='#334E68')
    im.save(path)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(total))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(widths[i]))
            tc_w.set(qn('w:type'), 'dxa')
            cell.width = Inches(widths[i] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def style_run(run, size=10.5, color='172B4D', bold=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn('w:ascii'), FONT)
    rpr.rFonts.set(qn('w:hAnsi'), FONT)
    lang = rpr.find(qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang')
        rpr.append(lang)
    lang.set(qn('w:val'), 'ko-KR')
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def configure_new_doc(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.82)
    sec.bottom_margin = Inches(0.82)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn('w:ascii'), FONT)
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in [
        ('Heading 1', 16, BLUE, 16, 8), ('Heading 2', 13, BLUE, 12, 6), ('Heading 3', 12, NAVY, 8, 4)
    ]:
        s = styles[name]
        s.font.name = FONT
        s._element.rPr.rFonts.set(qn('w:ascii'), FONT)
        s._element.rPr.rFonts.set(qn('w:hAnsi'), FONT)
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.font.bold = True
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(6)
    style_run(p.add_run(title), 25, NAVY, True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    style_run(p.add_run(subtitle), 12.5, GRAY)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(18)
    style_run(meta.add_run('MUSUBI  |  기준일 2026-08-18  |  상태: 최종 확정 목표'), 10, BLUE, True)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass
    for i, text in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], 'E8EEF5')
        p = table.rows[0].cells[i].paragraphs[0]
        style_run(p.add_run(text), 9.5, NAVY, True)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            style_run(p.add_run(str(text)), 9.2, '27364B')
    set_table_geometry(table, widths)
    return table


def add_heading_safe(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6)
    style_run(p.add_run(text), 16 if level == 1 else 13, BLUE, True)
    return p


def add_footer(doc, label):
    for sec in doc.sections:
        p = sec.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        style_run(p.add_run(label), 8.5, GRAY)


def new_korean_doc():
    """Start from the verified Korean reference theme, then remove its body."""
    template = ROOT / 'Infra/project-docs/reference/product/CineVerse_요구사항정의서_v3_final.docx'
    doc = Document(template)
    body = doc._element.body
    for child in list(body):
        if child.tag != qn('w:sectPr'):
            body.remove(child)
    for sec in doc.sections:
        for part in (sec.header, sec.footer):
            for p in part.paragraphs:
                p.clear()
    return doc


def build_architecture_doc(logical: Path, physical: Path):
    doc = new_korean_doc()
    configure_new_doc(doc)
    add_title(doc, 'Musubi 최종 클라우드 아키텍처', '2개 가용 영역 · 총 10 VM · 관리형 Load Balancer 기반 혼합형 구조')
    p = doc.add_paragraph()
    set_cell = doc.add_table(rows=1, cols=1)
    try:
        set_cell.style = 'Table Grid'
    except KeyError:
        pass
    cell = set_cell.cell(0, 0)
    set_cell_shading(cell, 'EAF2F8')
    cell.text = ''
    style_run(cell.paragraphs[0].add_run('확정 결론  현재 5 VM 운영 환경을 2개 AZ의 10 VM 구조로 확장한다. 프론트엔드와 백엔드는 KKE, AI와 DB는 전용 VM으로 운영한다.'), 11, NAVY, True)
    set_table_geometry(set_cell, [9360])

    doc.add_heading('1. 현재와 목표', level=1)
    add_table(doc, ['구분', '현재 운영', '최종 목표'], [
        ['KKE Worker', '2', '4 (AZ별 2)'], ['GPU AI', '1', '2 (AZ별 1)'],
        ['PostgreSQL', '1', '2 (Primary/Standby)'], ['NAT+Bastion', '1', '2 (AZ별 1)'], ['합계', '5', '10'],
    ], [2800, 2500, 4060])
    p = doc.add_paragraph('※ 관리형 LB, HA Group, KKE control plane, Container Registry, Object Storage는 VM 수에 포함하지 않는다.')
    for r in p.runs: style_run(r, 9, GRAY)

    doc.add_page_break()
    doc.add_heading('2. 논리 구조', level=1)
    doc.add_picture(str(logical), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('사용자 트래픽과 내부 AI 트래픽은 서로 다른 관리형 LB 계층으로 분리한다. 애플리케이션은 KKE, 상태 저장 DB와 GPU 실행환경은 전용 VM에 둔다.')
    for r in p.runs: style_run(r)

    doc.add_page_break()
    doc.add_heading('3. 물리 구조', level=1)
    doc.add_picture(str(physical), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table(doc, ['AZ', 'Public subnet', 'Private subnet'], [
        ['kr-central-2-a', '10.30.1.0/24 · NAT+Bastion-A', '10.30.2.0/24 · Worker-A1/A2, GPU-A, DB Primary'],
        ['kr-central-2-b', '10.30.4.0/24 · NAT+Bastion-B', '10.30.3.0/24 · Worker-B1/B2, GPU-B, DB Standby'],
    ], [2000, 2900, 4460])

    doc.add_page_break()
    doc.add_heading('4. 통신 및 라우팅', level=1)
    add_table(doc, ['구간', '경로/정책'], [
        ['외부 요청', 'movieverse.cloud → Public LB HA → LB-A/B → ingress-nginx → FE/BE Pods'],
        ['Ingress', 'HTTP NodePort 31664, HTTPS NodePort 31364, externalTrafficPolicy=Cluster'],
        ['AI 요청', 'Backend Pod → Internal AI LB HA → GPU-A/B → AI FastAPI → llama-server/Milvus'],
        ['DB 요청', 'Backend Pod → PgBouncer → PostgreSQL Primary'],
        ['DB 복제', 'PostgreSQL Primary → Standby streaming replication, 장애 시 수동 승격'],
        ['Public route', '0.0.0.0/0 → Internet Gateway'],
        ['Private-A route', '0.0.0.0/0 → NAT+Bastion-A'],
        ['Private-B route', '0.0.0.0/0 → NAT+Bastion-B'],
    ], [2500, 6860])

    doc.add_heading('5. 보안 원칙', level=1)
    add_table(doc, ['영역', '확정 정책'], [
        ['인터넷 공개', 'Public LB 80/443만 공개'],
        ['관리 접속', 'SSH는 AZ별 NAT+Bastion 경유'],
        ['내부 전용', 'KKE Worker, PostgreSQL, GPU AI VM'],
        ['차단 포트', '5432, 6432, 8081, 19530, 9000, 9001 인터넷 비공개'],
        ['권한 범위', 'DB/AI 인바운드는 필요한 내부 보안 그룹 간 통신만 허용'],
    ], [2500, 6860])

    doc.add_page_break()
    doc.add_heading('6. 고가용성과 복구', level=1)
    add_table(doc, ['구성요소', '대응 방식', '주의점'], [
        ['KKE', '4 Worker를 2 AZ에 분산, PDB/topology spread 적용', '동일 AZ 집중 배치 방지'],
        ['Public LB', 'LB-A/B + HA Group', '두 AZ 대상 헬스 체크'],
        ['GPU AI', 'Internal AI LB가 정상 GPU로 분산', '모델·런타임·설정·Milvus 정합성 유지'],
        ['PostgreSQL', 'Primary→Standby streaming replication', '자동 승격 아님; 운영자 수동 승격'],
        ['백업', 'Object Storage에 DB 백업', '복원 절차와 보존 정책 별도 검증'],
    ], [1900, 4000, 3460])

    doc.add_heading('7. 구축 완료 판정', level=1)
    checks = [
        '10 VM이 지정 AZ와 subnet에 배치됨', 'KKE Worker 4대 Ready 및 Pod AZ 분산 확인',
        'GPU-A/B 모델 체크섬·런타임·API 설정 일치', 'Internal AI LB 헬스 체크와 요청 분산 정상',
        'DB streaming replication과 복제 지연 확인', 'DB 수동 승격 및 애플리케이션 재연결 검증',
        'HTTPS·로그인·검색·추천·채팅·리뷰 스모크 테스트 통과', 'Object Storage 백업 생성 및 복원 절차 확인',
    ]
    add_table(doc, ['번호', '완료 조건'], [[i + 1, x] for i, x in enumerate(checks)], [1100, 8260])
    add_footer(doc, 'Musubi · 최종 클라우드 아키텍처 · 2026-08-18')
    path = OUT / 'Musubi_클라우드아키텍처_최종_20260818.docx'
    doc.save(path)
    return path


def replace_all(doc, replacements):
    def apply_to_paragraph(p):
        full = ''.join(r.text for r in p.runs)
        new = full
        for old, value in replacements.items():
            new = new.replace(old, value)
        if new != full:
            for r in p.runs: r.text = ''
            if p.runs:
                p.runs[0].text = new
            else:
                p.add_run(new)
    for p in doc.paragraphs: apply_to_paragraph(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs: apply_to_paragraph(p)


def append_final_section(base, out_name, kind, physical):
    target = OUT / out_name
    copy2(base, target)
    doc = Document(target)
    if kind == 'wbs':
        replace_all(doc, {
            'MySQL 스키마 구현 및 마이그레이션': 'PostgreSQL 스키마 구현 및 마이그레이션',
            'MySQL DB 스키마 설계 및 구현': 'PostgreSQL DB 스키마 설계 및 구현',
            'MySQL DB 구현 및 마이그레이션': 'PostgreSQL DB 구현 및 마이그레이션',
            'Redis 캐시 / 세션 / 랭킹 설계': 'PostgreSQL 기반 세션 / 랭킹 설계',
            'Redis 캐시·세션·랭킹 구현': 'PostgreSQL 기반 세션·랭킹 구현',
            'Redis 설정': '세션·랭킹 DB 설정',
        })
    doc.add_page_break()
    add_heading_safe(doc, '최종 인프라 확정 변경 (2026-08-18)', level=1)
    p = doc.add_paragraph('현재 5 VM 운영 환경과 최종 10 VM 목표를 구분하며, 다음 항목을 최종 확장 범위로 확정한다.')
    for r in p.runs: style_run(r)
    if kind == 'wbs':
        add_table(doc, ['ID', '작업', '담당', '산출물/완료 조건'], [
            ['9.1', 'Public-B subnet 및 AZ별 라우팅 구성', 'BE2', '10.30.4.0/24, Public/Private route 검증'],
            ['9.2', 'KKE Worker 4대 AZ 분산', 'BE2', '4 Nodes Ready, Pod topology 분산'],
            ['9.3', 'GPU-A/B 동일 런타임 구성', 'AI, BE2', '모델 SHA-256 및 API 스모크 일치'],
            ['9.4', 'Internal AI LB HA 구성', 'BE2', 'GPU 2대 Healthy 및 분산 확인'],
            ['9.5', 'PostgreSQL Standby 복제 구성', 'BE2', 'streaming replication 및 lag 확인'],
            ['9.6', 'DB 수동 승격 훈련', 'BE1, BE2', '승격·연결전환·복구 절차 기록'],
            ['9.7', 'Public LB HA 및 서비스 검증', 'ALL', 'HTTPS 및 핵심 기능 스모크 통과'],
            ['9.8', '문서·발표자료 최종화', 'ALL', 'WBS/요구사항/기획서/아키텍처 일치'],
        ], [900, 3600, 1400, 3460])
    elif kind == 'requirements':
        add_table(doc, ['ID', '항목', '요구사항'], [
            ['NFR-06', '가용성', 'KKE Worker를 2개 AZ에 분산하고 Public LB HA를 통해 단일 AZ 장애 시 서비스 진입 경로를 유지한다.'],
            ['NFR-07', '데이터 복구', 'PostgreSQL Primary의 변경 사항을 Standby로 스트리밍 복제하고, 장애 시 수동 승격 절차와 Object Storage 백업으로 복구한다.'],
            ['NFR-08', 'AI 가용성', '동일 운영 모델과 런타임의 GPU VM 2대를 Internal AI LB HA 뒤에 배치하고 정상 대상에 요청을 분산한다.'],
            ['NFR-09', '네트워크 보안', '외부에는 Public LB 80/443만 공개하며 KKE Worker, DB, GPU와 내부 서비스 포트는 Private network로 제한한다.'],
            ['NFR-10', '운영 검증', '배포 후 노드·Pod·LB·GPU·DB 복제·백업 상태와 핵심 사용자 기능을 스모크 테스트한다.'],
        ], [1400, 1800, 6160])
        add_heading_safe(doc, '인프라 제약 및 복구 정책', level=2)
        p = doc.add_paragraph('PostgreSQL Standby의 승격은 자동이 아니라 운영자 수동 절차로 수행한다. GPU-A/B의 모델, 런타임, AI API와 벡터 데이터는 동일 상태를 유지해야 하며 시험 모델은 운영 배포에서 제외한다.')
        for r in p.runs: style_run(r)
    elif kind == 'planning':
        doc.add_picture(str(physical), width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_table(doc, ['영역', '최종 기획'], [
            ['배치', '2개 AZ에 KKE 4, GPU 2, PostgreSQL 2, NAT+Bastion 2 - 총 10 VM'],
            ['사용자 트래픽', 'Public LB HA → ingress-nginx → Frontend/Backend Pods'],
            ['AI 트래픽', 'Backend → Internal AI LB HA → GPU-A/B'],
            ['데이터', 'PostgreSQL Primary/Standby streaming replication + Object Storage backup'],
            ['운영', '모니터링은 Kubernetes에 배치, 별도 Ops VM 없음'],
            ['복구', 'DB는 수동 승격, GPU는 정상 대상 라우팅, KKE는 AZ 분산'],
        ], [2200, 7160])
    add_footer(doc, f'Musubi · {kind} · 최종 인프라 변경 2026-08-18')
    doc.save(target)
    return target


def main():
    logical = ASSET / 'cineverse-multi-az-logical-20260818.png'
    physical = ASSET / 'cineverse-multi-az-physical-20260818.png'
    make_logical_diagram(logical)
    make_physical_diagram(physical)
    outputs = [build_architecture_doc(logical, physical)]
    outputs.append(append_final_section(ROOT / 'Infra/project-docs/reference/product/CineVerse_WBS.docx', 'Musubi_WBS_최종_20260818.docx', 'wbs', physical))
    outputs.append(append_final_section(ROOT / 'Infra/project-docs/reference/product/CineVerse_요구사항정의서_v3_final.docx', 'Musubi_요구사항정의서_최종_20260818.docx', 'requirements', physical))
    outputs.append(append_final_section(ROOT / 'Infra/project-docs/reference/product/CineVerse_기획서_v4_final.docx', 'Musubi_기획서_최종_20260818.docx', 'planning', physical))
    print('\n'.join(str(p) for p in outputs))


if __name__ == '__main__':
    main()
