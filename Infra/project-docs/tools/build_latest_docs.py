from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
OUT = ROOT / "Infra/project-docs/final-delivery/2026-08-17"
FONT = "Arial Unicode MS"
NAVY = "17324D"
BLUE = "2E74B5"
ORANGE = "F59E0B"
INK = "1F2937"
MUTED = "5B6573"
LIGHT = "F2F4F7"
PALE_BLUE = "EAF2F8"
PALE_ORANGE = "FFF4E0"
GREEN = "177245"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=80, start=110, bottom=80, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def apply_font_to_run(run):
    run.font.name = FONT
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), FONT)


def set_run(run, size=None, bold=False, color=INK):
    apply_font_to_run(run)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def iter_paragraphs(container):
    yield from container.paragraphs
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def apply_document_fonts(doc):
    for style in doc.styles:
        if not hasattr(style, "_element") or style._element.rPr is None:
            continue
        style.font.name = FONT
        r_fonts = style._element.rPr.get_or_add_rFonts()
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
            r_fonts.set(qn(f"w:{attribute}"), FONT)

    for paragraph in iter_paragraphs(doc):
        for run in paragraph.runs:
            apply_font_to_run(run)

    for section in doc.sections:
        for part in (section.header, section.footer):
            for paragraph in iter_paragraphs(part):
                for run in paragraph.runs:
                    apply_font_to_run(run)


def save_document(doc, filename):
    apply_document_fonts(doc)
    doc.save(OUT / filename)


def add_field(paragraph, field):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])
    set_run(run, 8, color=MUTED)


def setup_document(title):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.82)
    sec.right_margin = Inches(0.82)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in (
        ("Title", 28, NAVY, 0, 10),
        ("Heading 1", 17, NAVY, 13, 7),
        ("Heading 2", 12.5, BLUE, 9, 5),
        ("Heading 3", 10.5, NAVY, 6, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(9.5)
        style.paragraph_format.left_indent = Inches(0.26)
        style.paragraph_format.first_line_indent = Inches(-0.16)
        style.paragraph_format.space_after = Pt(3)

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"MUSUBI  |  {title}")
    set_run(run, 8, bold=True, color=MUTED)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("최종 검증 기준 2026-08-19  ·  ")
    set_run(run, 8, color=MUTED)
    add_field(footer, "PAGE")
    return doc


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text.upper())
    set_run(r, 9, bold=True, color=ORANGE)
    r.font.all_caps = True


def add_title_block(doc, title, subtitle, status):
    add_kicker(doc, "FINAL PRESENTATION BRIEF")
    doc.add_heading(title, 0)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(subtitle)
    set_run(r, 11, color=MUTED)
    add_callout(doc, status, PALE_ORANGE, ORANGE)


def add_callout(doc, text, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [180, 9180])
    shade(table.cell(0, 0), accent)
    shade(table.cell(0, 1), fill)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, 105, 125, 105, 125)
    p = table.cell(0, 1).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run(r, 9.5, bold=True, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        shade(cell, LIGHT)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run(r, font_size, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_margins(cells[idx])
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run(r, font_size, color=INK)
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run(r, 9.5, color=INK)


def add_page_break(doc):
    doc.add_page_break()


def build_cloud():
    doc = setup_document("클라우드 아키텍처")
    add_title_block(
        doc,
        "Musubi 클라우드 아키텍처",
        "KakaoCloud 멀티 AZ 운영 구성 · 구현 완료 기준",
        "현재 상태: 2026-08-19 기준 2개 가용 영역, 10개 VM, Public/AI 로드밸런서 고가용성, PostgreSQL Primary–Standby 구성이 완료되었습니다.",
    )
    doc.add_heading("1. 구성 원칙", level=1)
    add_bullets(doc, [
        "외부 진입점은 Public HA Group과 영역별 ALB로 이중화하고 HTTPS 트래픽을 KKE Ingress로 전달합니다.",
        "애플리케이션·GPU·DB는 Private Subnet에 두고, 관리 및 아웃바운드는 영역별 NAT+Bastion을 사용합니다.",
        "GPU API는 Internal AI ALB를 통해 두 GPU 노드로 분산하며 인터넷에 직접 노출하지 않습니다.",
        "PostgreSQL은 AZ-A Primary와 AZ-B Standby의 비동기 Streaming Replication으로 구성합니다.",
    ])
    doc.add_heading("2. 운영 자원 인벤토리", level=1)
    add_table(doc, ["계층", "AZ-A", "AZ-B", "합계"], [
        ("KKE Worker", "10.30.2.178 / 10.30.2.122", "10.30.3.111 / 10.30.3.66", "4"),
        ("GPU AI", "10.30.2.227", "10.30.3.119", "2"),
        ("PostgreSQL", "Primary 10.30.2.185", "Standby 10.30.3.190", "2"),
        ("NAT + Bastion", "10.30.1.134", "10.30.4.87", "2"),
        ("총 VM", "5", "5", "10"),
    ], [1800, 3300, 3300, 960], 8.2)
    doc.add_heading("3. 핵심 트래픽 경로", level=1)
    add_table(doc, ["구분", "경로"], [
        ("사용자", "DNS(movieverse.cloud) → Public HA Group → ALB A/B → KKE Ingress → Frontend/Backend"),
        ("AI 추론", "Backend Pod → Internal AI ALB A/B → GPU-A 또는 GPU-B :80"),
        ("데이터", "Backend Pod → PostgreSQL Primary; Primary → Standby 비동기 복제"),
        ("운영", "관리자 → Bastion A/B → Private VM; Private outbound → NAT A/B"),
    ], [1800, 7560], 8.2)

    add_page_break(doc)
    add_kicker(doc, "NETWORK & HIGH AVAILABILITY")
    doc.add_heading("네트워크와 고가용성", 0)
    doc.add_heading("1. VPC 및 서브넷", level=1)
    add_table(doc, ["영역", "서브넷", "CIDR", "주요 배치"], [
        ("AZ-A", "Public A", "10.30.1.0/24", "ALB-A, NAT+Bastion-A"),
        ("AZ-A", "Private A", "10.30.2.0/24", "KKE-A×2, GPU-A, DB Primary"),
        ("AZ-B", "Public B", "10.30.4.0/24", "ALB-B, NAT+Bastion-B"),
        ("AZ-B", "Private B", "10.30.3.0/24", "KKE-B×2, GPU-B, DB Standby"),
    ], [1300, 1900, 1900, 4260])
    doc.add_heading("2. 장애 격리 방식", level=1)
    add_bullets(doc, [
        "Public HA Group은 두 영역의 ALB 노드를 묶어 단일 서비스 진입점을 제공합니다.",
        "KKE Worker는 영역별 2대로 구성되어 Frontend·Backend·Ingress Pod를 분산 배치합니다.",
        "Internal AI ALB는 /health 응답을 기준으로 GPU API 대상의 상태를 확인합니다.",
        "DB Standby는 read-only이며, Primary 장애 시 운영자가 수동 승격하는 방식입니다.",
    ])
    doc.add_heading("3. 완료된 검증", level=1)
    add_table(doc, ["검증 항목", "결과"], [
        ("KKE", "Worker 4대 Ready, Pod 영역 분산, PDB·topology spread 적용"),
        ("외부 경로", "Public HA/DNS 2개 IP 및 Frontend/Backend HTTP 200 확인"),
        ("AI 경로", "Internal AI LB 대상 정상, GPU-A/B API health 확인"),
        ("DB 복제", "PostgreSQL 17.11 비동기 streaming, Standby read-only 확인"),
    ], [2500, 6860])
    add_callout(doc, "보안 경계: 외부에는 Public HA의 80/443만 노출하고 GPU·DB·KKE 노드는 Private Subnet에서 내부 통신만 허용합니다.")

    add_page_break(doc)
    add_kicker(doc, "OPERATIONS & REMAINING TESTS")
    doc.add_heading("운영 기준과 남은 검증", 0)
    doc.add_heading("1. 운영 전환 기준", level=1)
    add_bullets(doc, [
        "Public LB는 HTTP 80을 HTTPS 443으로 리다이렉트하고, TLS 인증서는 movieverse.cloud에 연결합니다.",
        "Internal AI LB의 DNS 이름을 Backend의 AI_BASE_URL로 사용합니다.",
        "DB 쓰기는 Primary에만 수행하고 Standby는 장애 복구·복제 검증 용도로 유지합니다.",
        "NAT+Bastion은 관리 접속과 Private 인스턴스 아웃바운드의 공용 역할을 수행합니다.",
    ])
    doc.add_heading("2. 확인된 제약", level=1)
    add_table(doc, ["항목", "현재 상태", "후속 조치"], [
        ("DB 전환", "수동 승격", "장애 훈련 후 자동화 여부 결정"),
        ("DB Standby 볼륨", "루트 50GB", "운영 장기화 시 100GB 이상 확장"),
        ("GPU 장애 전환", "LB 상태 확인 구성", "GPU-A 강제 중지 후 연속 요청 검증"),
        ("백업 복구", "Object Storage 사용", "복구 리허설과 RPO/RTO 확정"),
    ], [1800, 2500, 5060])
    doc.add_heading("3. 발표 시 핵심 설명", level=1)
    add_callout(doc, "Musubi는 확장 가능한 Stateless 계층은 Kubernetes로, 상태·가속기 의존성이 큰 PostgreSQL과 GPU AI는 전용 VM으로 분리한 혼합형 아키텍처입니다.", PALE_ORANGE, ORANGE)
    add_bullets(doc, [
        "단일 AZ·단일 GPU 구조에서 두 AZ·이중 GPU·DB Standby 구조로 확장했습니다.",
        "모든 계층을 Kubernetes에 넣지 않고 워크로드 특성에 따라 관리 경계를 분리했습니다.",
        "현재 구성 완료와 기능 검증은 구분해 기록했으며, 실제 장애 전환 훈련은 후속 운영 과제입니다.",
    ])
    save_document(doc, "Musubi_클라우드아키텍처_최신.docx")


def build_ai():
    doc = setup_document("AI 변경사항")
    add_title_block(
        doc,
        "Musubi AI 변경사항",
        "Gemma 4 12B 기반 대화·추천 파이프라인 · 운영 반영 기준",
        "현재 상태: 검증된 AI 런타임 10개 파일을 GPU-B → GPU-A 순으로 롤링 배포했고, 두 노드의 파일 해시·서비스 health·추천 스모크를 확인했습니다.",
    )
    doc.add_heading("1. 운영 모델과 실행 기준", level=1)
    add_table(doc, ["항목", "운영 기준"], [
        ("모델", "gemma-4-12b-it-base-q4_k_m.gguf (Gemma 4 12B, Q4_K_M)"),
        ("모델 SHA-256", "9808e158…d64bf1a4"),
        ("llama.cpp", "build 9776 / commit ac4105d…"),
        ("실행 옵션", "ctx-size 20480, np=5, reasoning-budget 0, GPU offload"),
        ("금지 옵션", "--skip-chat-parsing 사용 금지"),
        ("벡터 컬렉션", "movies_active / characters_verified_v5"),
    ], [2200, 7160], 8.5)
    doc.add_heading("2. 파이프라인 개선 범위", level=1)
    add_bullets(doc, [
        "캐릭터별 프로필·실용문장 fallback과 관계 근거를 분리해 말투 평준화와 근거 없는 관계 단정을 줄였습니다.",
        "비밀·물건·면접·발표 등 다중 턴 문맥과 안전 규칙을 강화하고, 공격적 표현과 상대 의도 단정을 차단했습니다.",
        "영화 추천은 구조화된 카드의 제목·장르·추천 이유만 사용해 설명과 실제 카드의 불일치를 줄였습니다.",
        "명확한 조건 검색은 CrossEncoder 생략 경로를 사용하고, 자연어·복합 검색은 재정렬 경로를 사용합니다.",
        "날짜·국가·장르·평점 등 명시 조건은 검색 단계의 hard filter로 처리하고, 확인할 수 없는 실시간 OTT 정보는 경계를 안내합니다.",
    ])
    add_callout(doc, "모델을 다시 학습하기보다 동일 베이스 모델에서 라우팅·검증·fallback·검색 정책을 개선해 품질과 롤백 안정성을 확보했습니다.")

    add_page_break(doc)
    add_kicker(doc, "VERIFICATION & PRODUCTION RELEASE")
    doc.add_heading("검증 결과와 운영 배포", 0)
    doc.add_heading("1. 누적 품질 검증", level=1)
    add_table(doc, ["검증", "결과"], [
        ("캐릭터 50명 × 3상황 × 3턴", "450/450 통과"),
        ("유사도 0.90 이상 답변", "86쌍 → 0쌍"),
        ("실사용자형 강건성", "15/15 통과"),
        ("배포 전 회귀", "255/255 통과"),
        ("배포 전 실제 API 스모크", "7/7 통과 / critical failure 0"),
        ("최종 로컬 검증", "432 tests + 108 subtests 통과"),
    ], [5200, 4160])
    doc.add_heading("2. 2026-08-19 운영 반영", level=1)
    add_bullets(doc, [
        "배포 순서: GPU-B(10.30.3.119) → GPU-A(10.30.2.227).",
        "범위: pipeline/, rag/, data/topic_profiles.json 아래 승인된 런타임·데이터 파일 10개.",
        "제외: GGUF, llama-server 옵션, .env, 가상환경, Milvus 데이터, 로그, 테스트 모델.",
        "롤백 사본: 각 노드 /home/ubuntu/cineverse-backups/ai-final-20260819-101047/.",
        "배포 후 두 노드의 SHA-256 일치와 cineverse-api.service health를 확인했습니다.",
    ])
    doc.add_heading("3. 운영 스모크", level=1)
    add_table(doc, ["대상", "검증 결과"], [
        ("GPU-A", "HTTP 200, 영화 카드 3개, 단일 표본 15.77초"),
        ("GPU-B", "HTTP 200, 영화 카드 3개, 단일 표본 4.87초"),
        ("Internal AI HA", "health 0.169초 / 추천 11.87초 / 카드 3개"),
    ], [2600, 6760])
    add_callout(doc, "GPU-A/B 시간은 각각 한 번의 배포 확인 표본이며 성능 우열을 단정하는 벤치마크가 아닙니다.", PALE_ORANGE, ORANGE)

    add_page_break(doc)
    add_kicker(doc, "PERFORMANCE & OPERATING BOUNDARIES")
    doc.add_heading("성능과 운영 경계", 0)
    doc.add_heading("1. 확인된 성능 개선", level=1)
    add_table(doc, ["지표", "변경 전", "변경 후", "개선"], [
        ("단일 추천", "7.719초", "6.494초", "15.9%"),
        ("5개 동시 요청 p95", "21.555초", "19.103초", "11.4%"),
        ("조건 검색 생략 경로", "-", "0.273초", "CrossEncoder 생략"),
    ], [2500, 1900, 1900, 3060])
    doc.add_heading("2. 현재 한계", level=1)
    add_bullets(doc, [
        "단일 T4 노드에서 동시 생성 시 요청별 token/s가 5.27~9.33까지 낮아진 측정 이력이 있습니다.",
        "10개 동시 추천 요청 p95는 41.520초였고, 20개 이상에서는 admission 한도에 따라 429/503이 발생했습니다.",
        "캐릭터 원작 지식은 검증된 내부 데이터 범위까지만 답할 수 있습니다.",
        "실시간 OTT 제공 여부는 현재 컬렉션에 지역·시점별 근거가 없어 확정 답변하지 않습니다.",
    ])
    doc.add_heading("3. 남은 운영 확인", level=1)
    add_table(doc, ["항목", "상태"], [
        ("Kubernetes AI_BASE_URL", "로컬 manifest는 Internal HA DNS를 가리킴"),
        ("Live cluster 적용", "kic-iam-auth 부재로 해당 배포 세션에서는 미확인"),
        ("장기 트래픽 A/B", "미수행; 운영 로그 축적 후 재평가"),
        ("GPU failover", "LB 구성 완료; 강제 장애 전환 시나리오 후속 검증"),
    ], [3400, 5960])
    add_callout(doc, "발표 시에는 ‘모든 문제가 해결됐다’가 아니라, 검증된 품질 개선·이중 GPU 배포·남은 운영 확인 항목을 분리해 설명합니다.")
    save_document(doc, "Musubi_AI_변경사항_최신.docx")


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    build_cloud()
    build_ai()
